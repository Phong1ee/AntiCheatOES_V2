"""Fills the shipped Question Bank import template with a chosen Subject.

The teacher picks a subject they are assigned to and gets the template back with
the Subject block already correct, so the Subject ID can no longer disagree with
the subject the request is filed against - the mismatch the upload rejects.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "assets" / "question-bank-import-template.docx"

SUBJECT_FIELDS = ("Subject ID", "Subject Name", "Description")


def _rewrite(paragraph: Paragraph, line: str) -> None:
    """Replace a paragraph's text, keeping the first run's formatting.

    The line lives in one run today, but Word re-splits runs whenever the asset
    is edited, so never assume a one-to-one mapping.
    """
    paragraph.runs[0].text = line
    for run in paragraph.runs[1:]:
        run.text = ""


def build_subject_template(subject_id: str, subject_name: str, description: str) -> bytes:
    """Return the template as DOCX bytes, with the sample Subject block filled in."""
    document = Document(str(TEMPLATE_PATH))
    values = {
        "Subject ID": subject_id.strip(),
        "Subject Name": subject_name.strip(),
        # The parser rejects an empty Description, so keep the prompt when the
        # subject has none rather than shipping a file that fails its own rules.
        "Description": description.strip() or "Brief subject description",
    }
    filled: set[str] = set()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for field in SUBJECT_FIELDS:
            # Only the sample block has "<field>: <value>" lines; the instructions
            # above it name the fields in prose, never followed by a colon.
            if field in filled or not paragraph.runs or not text.startswith(f"{field}:"):
                continue
            _rewrite(paragraph, f"{field}: {values[field]}")
            filled.add(field)
            break
    missing = [field for field in SUBJECT_FIELDS if field not in filled]
    if missing:
        raise RuntimeError(f"Import template is missing its {', '.join(missing)} line")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_subject_guideline(
    subject_id: str,
    subject_name: str,
    description: str,
    chapters: list[tuple[str, list[str]]],
) -> bytes:
    """A reference sheet of a subject's existing Chapters and Learning Objectives.

    Import matches these by name: an exact existing name is reused, anything else
    creates a new Chapter or Learning Objective. Teachers cannot see the taxonomy
    while filling the template, so they invent near-miss names and silently grow
    duplicates. This is deliberately a separate document - appending it to the
    template would put "CHAPTER:" lines in front of the parser.
    """
    document = Document()
    document.add_heading("Question Bank Guideline", level=0)
    document.add_paragraph(f"{subject_id} - {subject_name}")
    if description.strip():
        document.add_paragraph(description.strip())

    document.add_heading("How to use this", level=1)
    document.add_paragraph(
        "Copy a Chapter or Learning Objective name below exactly as written to add "
        "questions to it. Any other name creates a new one during import.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Names are matched ignoring case and surrounding spaces. Learning Objectives "
        "belong to a Chapter, so the same name under another Chapter is a different one.",
        style="List Bullet",
    )
    document.add_paragraph(
        "In the template, a question's Learning Objectives line separates names with |.",
        style="List Bullet",
    )

    document.add_heading("Existing Chapters and Learning Objectives", level=1)
    if not chapters:
        document.add_paragraph(
            "This subject has no chapters yet. Every Chapter and Learning Objective "
            "in your file will be created during import."
        )
    seen: dict[str, int] = {}
    for chapter_name, lo_names in chapters:
        key = chapter_name.strip().casefold()
        seen[key] = seen.get(key, 0) + 1
        document.add_heading(chapter_name, level=2)
        if not lo_names:
            document.add_paragraph("No Learning Objectives yet.")
        for lo_name in lo_names:
            document.add_paragraph(lo_name, style="List Bullet")

    duplicates = sorted({name for name, _ in chapters if seen.get(name.strip().casefold(), 0) > 1})
    if duplicates:
        # Ambiguity is what makes an import fail outright, so name it here.
        document.add_heading("Needs an administrator", level=1)
        document.add_paragraph(
            "These Chapter names appear more than once, so import cannot tell them "
            "apart and will reject questions using them: " + ", ".join(duplicates)
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
