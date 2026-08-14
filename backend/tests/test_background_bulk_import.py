import os
import tempfile
import unittest
from datetime import date
from io import BytesIO
from unittest.mock import patch

from fastapi import HTTPException
from docx import Document
from openpyxl import Workbook
from sqlalchemy import create_engine, event
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from database import Base
from src.a_db_config import BackgroundJob, BackgroundJobStatus, BackgroundJobType, OutboxEvent, Question, Subject, User, UserRole
from src.route.adminRoute import _owned_import_job, import_users, import_users_from_rows
from src.service.import_job_service import import_job_summary, queue_import_job
from src.service.import_worker import handle_import_requested
from src.service.user_import_service import parse_user_import_xlsx


HEADERS = ["school_id", "full_name", "email", "role", "phone", "date_of_birth", "initial_password"]


def workbook_bytes(rows):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Users"
    sheet.append(HEADERS)
    for row in rows:
        sheet.append(row)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def question_document_bytes(count=1):
    document = Document()
    for line in (
        "SUBJECT", "Subject ID: DB", "Subject Name: Databases", "Description: Database concepts.",
        "CHAPTER: Transactions",
    ):
        document.add_paragraph(line)
    for number in range(1, count + 1):
        for line in (
            f"QUESTION {number}", "Type: Multiple Choice", "Difficulty: Easy",
            "Learning Objectives: Explain ACID", f"Content: Which property preserves all-or-nothing work? {number}",
            "A. Atomicity", "B. Duplication", "Answer: A",
        ):
            document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class BackgroundBulkImportTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
        event.listen(cls.engine, "connect", lambda connection, _: connection.execute("PRAGMA foreign_keys=ON"))
        cls.Session = sessionmaker(bind=cls.engine, expire_on_commit=False)

    def setUp(self):
        Base.metadata.drop_all(self.engine)
        Base.metadata.create_all(self.engine)
        self.db = self.Session()
        self.db.add_all([
            User(school_id="A1", full_name="Admin", email="admin@bulk.test", password_hash="x", role=UserRole.admin),
            User(school_id="A2", full_name="Other Admin", email="other@bulk.test", password_hash="x", role=UserRole.admin),
        ])
        self.db.commit()
        self.temp = tempfile.TemporaryDirectory()
        self.environment = patch.dict(os.environ, {"IMPORT_STAGING_DIR": self.temp.name}, clear=False)
        self.environment.start()

    def tearDown(self):
        self.environment.stop()
        self.temp.cleanup()
        self.db.close()

    @staticmethod
    def _rows(count, *, invalid=False):
        rows = [
            [f"S{index}", f"Student {index}", f"student{index}@bulk.test", "student", None, date(2005, 1, 1), "password123"]
            for index in range(1, count + 1)
        ]
        if invalid:
            rows[-1][2] = rows[0][2]
        return rows

    def _queue(self, content, total_rows):
        job, duplicate = queue_import_job(
            self.db,
            job_type=BackgroundJobType.user_import,
            requested_by="A1",
            filename="users.xlsx",
            content=content,
            total_rows=total_rows,
            scope="users",
        )
        self.db.commit()
        return job, duplicate

    def test_small_sync_import_remains_all_or_nothing(self):
        result = import_users_from_rows(parse_user_import_xlsx(workbook_bytes(self._rows(2))), self.db)
        self.assertEqual(result["imported_count"], 2)
        self.assertEqual(self.db.query(BackgroundJob).count(), 0)

    def test_large_import_queues_compact_outbox_then_worker_completes_once(self):
        content = workbook_bytes(self._rows(3))
        job, duplicate = self._queue(content, 3)
        self.assertFalse(duplicate)
        event_row = self.db.query(OutboxEvent).one()
        self.assertEqual(event_row.event_type, "import.requested")
        self.assertEqual(event_row.payload_json, {"job_id": job.job_id, "job_type": "USER_IMPORT"})
        self.assertNotIn("PK", str(event_row.payload_json))
        self.assertEqual(self.db.get(BackgroundJob, job.job_id).status, BackgroundJobStatus.pending)

        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)
        self.db.commit()
        completed = self.db.get(BackgroundJob, job.job_id)
        self.assertEqual(completed.status, BackgroundJobStatus.completed)
        self.assertEqual((completed.total_rows, completed.processed_rows, completed.success_rows, completed.failed_rows), (3, 3, 3, 0))
        self.assertEqual(self.db.query(User).filter(User.school_id.like("S%")).count(), 3)
        self.assertFalse(os.path.exists(os.path.join(self.temp.name, f"import_job_{job.job_id}.xlsx")))

        # A restart/redelivery after completion does not run the importer twice.
        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)
        self.db.commit()
        self.assertEqual(self.db.query(User).filter(User.school_id.like("S%")).count(), 3)

    def test_invalid_large_import_fails_without_creating_any_users(self):
        job, _ = self._queue(workbook_bytes(self._rows(2, invalid=True)), 2)
        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)
        self.db.commit()
        failed = self.db.get(BackgroundJob, job.job_id)
        self.assertEqual(failed.status, BackgroundJobStatus.failed)
        self.assertEqual(self.db.query(User).filter(User.school_id.like("S%")).count(), 0)

    def test_same_source_uses_one_job_and_status_hides_source_details(self):
        content = workbook_bytes(self._rows(2))
        first, duplicate = self._queue(content, 2)
        second, duplicate_retry = self._queue(content, 2)
        self.assertFalse(duplicate)
        self.assertTrue(duplicate_retry)
        self.assertEqual(first.job_id, second.job_id)
        summary = import_job_summary(first)
        self.assertNotIn("source_sha256", str(summary))
        self.assertNotIn("source_filename", str(summary))

        with self.assertRaises(HTTPException) as denied:
            _owned_import_job(self.db, first.job_id, "A2")
        self.assertEqual(denied.exception.status_code, 404)

    def test_large_api_returns_pending_job_before_worker_and_admin_job_access_is_owned(self):
        from fastapi import UploadFile

        content = workbook_bytes(self._rows(3))
        with patch("src.route.adminRoute.should_background_import", return_value=True):
            result = __import__("asyncio").run(
                import_users(UploadFile(filename="users.xlsx", file=BytesIO(content)), {"school_id": "A1"}, {}, self.db)
            )
        self.assertTrue(result["background"])
        self.assertEqual(result["status"], "PENDING")
        self.assertEqual(self.db.query(User).filter(User.school_id.like("S%")).count(), 0)
        self.assertEqual(self.db.get(BackgroundJob, result["jobId"]).requested_by, "A1")

    def test_question_import_reuses_existing_taxonomy_importer_in_background(self):
        self.db.add(Subject(subject_id="DB", subject_name="Databases", subject_description="Database concepts"))
        self.db.commit()
        content = question_document_bytes()
        job, _ = queue_import_job(
            self.db,
            job_type=BackgroundJobType.question_import,
            requested_by="A1",
            filename="questions.docx",
            content=content,
            total_rows=1,
            scope="subject:DB",
            metadata={"subject_id": "DB", "new_subject": False},
        )
        self.db.commit()
        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)
        self.db.commit()
        completed = self.db.get(BackgroundJob, job.job_id)
        self.assertEqual(completed.status, BackgroundJobStatus.completed)
        self.assertEqual(completed.success_rows, 1)

    def test_question_import_is_one_transaction_and_redelivery_does_not_duplicate(self):
        self.db.add(Subject(subject_id="DB", subject_name="Databases", subject_description="Database concepts"))
        self.db.commit()
        job, _ = queue_import_job(
            self.db,
            job_type=BackgroundJobType.question_import,
            requested_by="A1",
            filename="questions.docx",
            content=question_document_bytes(120),
            total_rows=120,
            scope="subject:DB:batched",
            metadata={"subject_id": "DB", "new_subject": False},
        )
        self.db.commit()
        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)

        completed = self.db.get(BackgroundJob, job.job_id)
        self.assertEqual((completed.processed_rows, completed.success_rows, completed.failed_rows), (120, 120, 0))
        self.assertEqual(completed.result_metadata["duplicate_skipped_count"], 0)

        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)
        self.db.commit()
        self.assertEqual(completed.success_rows, 120)

    def test_user_import_rejects_invalid_file_without_partial_rows(self):
        content = workbook_bytes(self._rows(5, invalid=True))
        job, _ = self._queue(content, 5)
        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)

        completed = self.db.get(BackgroundJob, job.job_id)
        self.assertEqual(completed.status, BackgroundJobStatus.failed)
        self.assertEqual(self.db.query(User).filter(User.school_id.like("S%")).count(), 0)

    def test_user_import_is_one_transaction(self):
        job, _ = self._queue(workbook_bytes(self._rows(120)), 120)
        handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)

        completed = self.db.get(BackgroundJob, job.job_id)
        self.assertEqual((completed.processed_rows, completed.success_rows, completed.failed_rows), (120, 120, 0))

    def test_fault_before_terminal_commit_rolls_back_question_rows(self):
        from src.service import import_worker

        self.db.add(Subject(subject_id="DB", subject_name="Databases", subject_description="Database concepts"))
        self.db.commit()
        job, _ = queue_import_job(
            self.db,
            job_type=BackgroundJobType.question_import,
            requested_by="A1",
            filename="questions.docx",
            content=question_document_bytes(2),
            total_rows=2,
            scope="subject:DB:fault",
            metadata={"subject_id": "DB", "new_subject": False},
        )
        self.db.commit()

        with patch.object(import_worker, "complete_job", side_effect=RuntimeError("forced pre-commit failure")):
            with self.assertRaisesRegex(RuntimeError, "forced pre-commit failure"):
                handle_import_requested({"event_type": "import.requested", "aggregate_id": str(job.job_id)}, self.db)
        self.db.rollback()

        self.assertEqual(self.db.query(Question).count(), 0)
        self.assertEqual(self.db.get(BackgroundJob, job.job_id).status, BackgroundJobStatus.running)
        self.assertTrue(os.path.exists(os.path.join(self.temp.name, f"import_job_{job.job_id}.docx")))


if __name__ == "__main__":
    unittest.main()
